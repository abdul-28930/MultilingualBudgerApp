import io
import csv
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract
from tabulate import tabulate


class DocumentAnalyzer:
    """Comprehensive document analyzer for financial documents."""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': 'PDF Document',
            '.doc': 'Word Document',
            '.docx': 'Word Document',
            '.xlsx': 'Excel Spreadsheet',
            '.xls': 'Excel Spreadsheet',
            '.csv': 'CSV File',
            '.png': 'Image',
            '.jpg': 'Image',
            '.jpeg': 'Image',
            '.bmp': 'Image',
            '.tiff': 'Image'
        }
    
    def get_file_type(self, file_path: Path) -> str:
        """Get the type of file based on extension."""
        suffix = file_path.suffix.lower()
        return self.supported_types.get(suffix, 'Unknown')
    
    async def analyze_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze document and return structured information."""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return await self._analyze_pdf(file_path)
        elif suffix in ['.doc', '.docx']:
            return await self._analyze_word_document(file_path)
        elif suffix in ['.xlsx', '.xls']:
            return await self._analyze_excel_file(file_path)
        elif suffix == '.csv':
            return await self._analyze_csv_file(file_path)
        elif suffix in {'.png', '.jpg', '.jpeg', '.bmp', '.tiff'}:
            return await self._analyze_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    async def _analyze_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Analyze PDF document."""
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            page_count = len(reader.pages)
            
            for page in reader.pages:
                page_text: Optional[str] = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            
            return {
                'file_type': 'PDF Document',
                'page_count': page_count,
                'text_content': full_text,
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'analysis_type': 'text_extraction',
                'summary': f"PDF document with {page_count} pages containing {len(full_text.split())} words"
            }
        except Exception as e:
            raise ValueError(f"Error analyzing PDF: {str(e)}")
    
    async def _analyze_word_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Word document."""
        try:
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    text_parts.append(" | ".join(row_data))
            
            full_text = "\n".join(text_parts)
            
            return {
                'file_type': 'Word Document',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'text_content': full_text,
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'analysis_type': 'text_extraction',
                'tables': table_data if table_data else None,
                'summary': f"Word document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
            }
        except Exception as e:
            raise ValueError(f"Error analyzing Word document: {str(e)}")
    
    async def _analyze_excel_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Excel file."""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            analysis = {
                'file_type': 'Excel Spreadsheet',
                'sheet_count': len(excel_data),
                'sheets': {},
                'analysis_type': 'data_analysis',
                'summary': f"Excel file with {len(excel_data)} sheets"
            }
            
            text_content = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name, df in excel_data.items():
                sheet_analysis = self._analyze_dataframe(df, sheet_name)
                analysis['sheets'][sheet_name] = sheet_analysis
                
                # Add to text content
                text_content.append(f"\n--- Sheet: {sheet_name} ---")
                text_content.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                text_content.append(f"Columns: {', '.join(df.columns.astype(str))}")
                
                # Add sample data
                if not df.empty:
                    text_content.append("Sample data:")
                    text_content.append(tabulate(df.head(5), headers='keys', tablefmt='pipe'))
                
                total_rows += len(df)
                total_cols += len(df.columns)
            
            analysis['text_content'] = "\n".join(text_content)
            analysis['total_rows'] = total_rows
            analysis['total_columns'] = total_cols
            analysis['summary'] = f"Excel file with {len(excel_data)} sheets, {total_rows} total rows"
            
            return analysis
            
        except Exception as e:
            raise ValueError(f"Error analyzing Excel file: {str(e)}")
    
    async def _analyze_csv_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze CSV file."""
        try:
            # Try to read with pandas
            df = pd.read_csv(file_path)
            
            analysis = self._analyze_dataframe(df, "CSV Data")
            
            # Add CSV-specific information
            analysis.update({
                'file_type': 'CSV File',
                'analysis_type': 'data_analysis',
                'summary': f"CSV file with {len(df)} rows and {len(df.columns)} columns"
            })
            
            # Create text content
            text_content = []
            text_content.append(f"CSV File Analysis")
            text_content.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            text_content.append(f"Columns: {', '.join(df.columns.astype(str))}")
            
            # Add sample data
            if not df.empty:
                text_content.append("\nSample data:")
                text_content.append(tabulate(df.head(10), headers='keys', tablefmt='pipe'))
            
            # Add statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_content.append("\nNumeric columns summary:")
                text_content.append(tabulate(df[numeric_cols].describe(), headers='keys', tablefmt='pipe'))
            
            analysis['text_content'] = "\n".join(text_content)
            
            return analysis
            
        except Exception as e:
            raise ValueError(f"Error analyzing CSV file: {str(e)}")
    
    async def _analyze_image(self, file_path: Path) -> Dict[str, Any]:
        """Analyze image using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                'file_type': 'Image',
                'image_size': image.size,
                'image_mode': image.mode,
                'text_content': text,
                'word_count': len(text.split()),
                'char_count': len(text),
                'analysis_type': 'ocr_extraction',
                'summary': f"Image ({image.size[0]}x{image.size[1]}) with {len(text.split())} words extracted via OCR"
            }
        except Exception as e:
            raise ValueError(f"Error analyzing image: {str(e)}")
    
    def _analyze_dataframe(self, df: pd.DataFrame, data_name: str) -> Dict[str, Any]:
        """Analyze a pandas DataFrame."""
        analysis = {
            'name': data_name,
            'rows': len(df),
            'columns': len(df.columns),
            'column_names': df.columns.tolist(),
            'data_types': df.dtypes.astype(str).to_dict(),
            'null_counts': df.isnull().sum().to_dict(),
            'numeric_columns': df.select_dtypes(include=['number']).columns.tolist(),
            'text_columns': df.select_dtypes(include=['object']).columns.tolist(),
            'date_columns': df.select_dtypes(include=['datetime']).columns.tolist(),
        }
        
        # Add statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            analysis['numeric_statistics'] = df[numeric_cols].describe().to_dict()
        
        # Identify potential financial columns
        financial_keywords = ['amount', 'price', 'cost', 'fee', 'balance', 'total', 'sum', 'revenue', 'income', 'expense']
        financial_cols = [col for col in df.columns if any(keyword in col.lower() for keyword in financial_keywords)]
        if financial_cols:
            analysis['potential_financial_columns'] = financial_cols
        
        return analysis


# Legacy function for backward compatibility
async def extract_text(file_path: Path) -> str:
    """Extract text from various file types - legacy function."""
    analyzer = DocumentAnalyzer()
    result = await analyzer.analyze_document(file_path)
    return result.get('text_content', '') 